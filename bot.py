import logging
import asyncio
import sqlite3
import re
import hashlib
import os
from docx import Document
from aiogram import Bot, Dispatcher, types, F
from aiogram.filters import Command
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiohttp import web

logging.basicConfig(level=logging.INFO)

TOKEN = os.getenv("BOT_TOKEN")
bot = Bot(token=TOKEN)
dp = Dispatcher()

async def handle(request):
    return web.Response(text="Bot is running smoothly!")

async def start_web_server():
    app = web.Application()
    app.router.add_get('/', handle)
    runner = web.AppRunner(app)
    await runner.setup()
    port = int(os.getenv("PORT", 10000))
    site = web.TCPSite(runner, '0.0.0.0', port)
    await site.start()

def init_db():
    conn = sqlite3.connect("quiz.db")
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS tests 
                      (id INTEGER PRIMARY KEY AUTOINCREMENT, subject TEXT, question TEXT, options TEXT, correct_index INTEGER)''')
    conn.commit()
    conn.close()

def get_subject_id(subject_name):
    return hashlib.md5(str(subject_name).encode('utf-8')).hexdigest()[:10]

def import_from_word():
    conn = sqlite3.connect("quiz.db")
    cursor = conn.cursor()
    cursor.execute("DELETE FROM tests")
    
    try:
        if not os.path.exists("testlar.docx"):
            print("Xatolik: testlar.docx topilmadi!")
            return
            
        doc = Document("testlar.docx")
        full_text = []
        
        # Word ichidagi barcha matn va jadvallarni bitta qatorga yig'amiz
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() and cell.text.strip() not in full_text:
                        full_text.append(cell.text.strip())
                        
        combined_text = "\n".join(full_text)
        
        # Hujjatni Mavzu: so'zi bo'yicha katta bloklarga ajratamiz
        subject_blocks = re.split(r'(?i)Mavzu:\s*', combined_text)
        
        for s_block in subject_blocks:
            s_block = s_block.strip()
            if not s_block:
                continue
                
            lines = s_block.split("\n")
            if lines and not re.match(r'^\d+[\.\s\)]', lines) and not lines.strip().startswith(("A)", "B)", "C)", "D)", "A.", "B.", "C.", "D.")):
                current_subject = lines.strip()
                block_content = "\n".join(lines[1:])
            else:
                current_subject = "Umumiy testlar"
                block_content = "\n".join(lines)
                
            # Savollarni Javob: qatoriga qarab xavfsiz ajratamiz (yonma-yon yozilgan matnlar uchun eng zo'r usul)
            raw_questions = re.split(r'(?i)(?:Javob:|Жавоб:|Тўғри жавоб:|To\'g\'ri javob:)\s*[A-DXZa-dxz]', block_content)
            answers_found = re.findall(r'(?i)(?:Javob:|Жавоб:|Тўғри жавоб:|To\'g\'ri javob:)\s*([A-DXZa-dxz])', block_content)
            
            for i, q_raw in enumerate(raw_questions):
                if i >= len(answers_found):
                    break
                    
                text_to_parse = q_raw.strip()
                if not text_to_parse:
                    continue
                
                # Savol matnini A variantgacha bo'lgan qismini qirqib olamiz
                q_match = re.search(r'^(?:\d+[\.\s\)]\s*)?(.*?)(?=[A-DXZa-dxz][\)\.])', text_to_parse, re.DOTALL)
                if not q_match:
                    continue
                question_body = q_match.group(1).strip()
                
                # Variantlarni bitta qatordan alohida qilib qidirib topamiz
                options = re.findall(r'(?:[A-DXZa-dxz][\)\.]\s*)(.*?)(?=[A-DXZa-dxz][\)\.]|$)', text_to_parse, re.DOTALL)
                clean_options = [o.strip() for o in options if o.strip()]
                
                if len(clean_options) > 4:
                    clean_options = clean_options[:4]
                    
                if len(clean_options) >= 2:
                    ans_letter = answers_found[i].upper()
                    correct_index = ord(ans_letter) - ord('A')
                    
                    if 0 <= correct_index < len(clean_options):
                        cursor.execute("INSERT INTO tests (subject, question, options, correct_index) VALUES (?, ?, ?, ?)",
                                       (current_subject, question_body, "&&".join(clean_options), correct_index))
                        conn.commit()
                        
        print("Word fayli muvaffaqiyatli o'qildi va xatolar tuzatildi!")
    except Exception as e:
        print("Word faylini o'qishda xatolik:", e)
    finally:
        conn.close()

@dp.message(Command("start"))
async def start_cmd(message: types.Message):
    conn = sqlite3.connect("quiz.db")
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT subject FROM tests")
    subjects = cursor.fetchall()
    conn.close()
    
    if not subjects:
        await message.answer("Bazada testlar hali ham topilmadi. Iltimos, Word faylingiz nomini testlar.docx ekanligini tekshiring.")
        return
        
    builder = InlineKeyboardBuilder()
    for sub in subjects:
        sub_name = sub
        builder.button(text=f"📂 {sub_name}", callback_data=f"sub_{get_subject_id(sub_name)}")
    builder.adjust(1)
    await message.answer("Assalomu alaykum! Mavzulardan birini tanlang va haqiqiy Telegram Quiz testini boshlang:", reply_markup=builder.as_markup())

@dp.callback_query(F.data.startswith("sub_") | F.data.startswith("nxt_"))
async def send_telegram_quiz(callback: types.CallbackQuery):
    target_sub_id = callback.data.split("_")
    conn = sqlite3.connect("quiz.db")
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT subject FROM tests")
    all_subjects = cursor.fetchall()
    
    subject_name = None
    for sub in all_subjects:
        if get_subject_id(sub) == target_sub_id:
            subject_name = sub
            break
            
    if not subject_name:
        await callback.message.answer("Mavzu topilmadi.")
        conn.close()
        return
        
    cursor.execute("SELECT question, options, correct_index FROM tests WHERE subject=? ORDER BY RANDOM() LIMIT 1", (subject_name,))
    question_data = cursor.fetchone()
    conn.close()
    
    if question_data:
        question, options, correct_index = question_data
        options_list = options.split("&&")
        await callback.message.answer_poll(
            question=f"📋 Mavzu: {subject_name}\n\n{question}"[:250],
            options=[opt[:100] for opt in options_list],
            type="quiz",
            correct_option_id=int(correct_index),
            is_anonymous=False,
            open_period=30
        )
        builder = InlineKeyboardBuilder()
        builder.button(text="Keyingi savol ➡️", callback_data=f"nxt_{target_sub_id}")
        builder.button(text="Menu 🏠", callback_data="back_to_menu")
        builder.adjust(1)
        await callback.message.answer("Savolga javob berib, keyingisiga o'tish tugmasini bosing:", reply_markup=builder.as_markup())
        try: await callback.message.delete()
        except: pass
    else:
        await callback.message.answer(f"📋 {subject_name} mavzusida testlar tugadi. Yangi mavzu uchun /start bosing.")

@dp.callback_query(F.data == "back_to_menu")
async def back_to_menu(callback: types.CallbackQuery):
    conn = sqlite3.connect("quiz.db")
    cursor = conn.cursor()
    cursor.execute("SELECT DISTINCT subject FROM tests")
    subjects = cursor.fetchall()
    conn.close()
    builder = InlineKeyboardBuilder()
    for sub in subjects:
        sub_name = sub
        builder.button(text=f"📂 {sub_name}", callback_data=f"sub_{get_subject_id(sub_name)}")
    builder.adjust(1)
    await callback.message.answer("Iltimos, test yechmoqchi bo'lgan **mavzuni tanlang**:", reply_markup=builder.as_markup())
    try: await callback.message.delete()
    except: pass

async def main():
    init_db()
    import_from_word()
    await asyncio.gather(start_web_server(), dp.start_polling(bot))

if __name__ == "__main__":
    asyncio.run(main())
